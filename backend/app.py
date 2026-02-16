from flask import Flask, request, jsonify
from flask_cors import CORS
import json
import os
import sys
import PyPDF2
import docx
import re

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from models.predict import ResumePredictor

app = Flask(__name__)
CORS(app)

# Configure upload settings
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max file size
app.config['UPLOAD_EXTENSIONS'] = ['.pdf', '.docx']

# Load predictor
predictor = ResumePredictor()

# Load interview questions
questions_path = os.path.join(os.path.dirname(__file__), 'data', 'interview_questions.json')
with open(questions_path, 'r') as f:
    interview_questions = json.load(f)


def extract_text_from_pdf(file):
    """Extract text from PDF file with better error handling"""
    try:
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        
        total_pages = len(pdf_reader.pages)
        # For very large PDFs, limit to first 50 pages to avoid memory issues
        max_pages = min(total_pages, 50)
        
        if total_pages > 50:
            print(f"Warning: PDF has {total_pages} pages. Processing first 50 pages only.")
        
        # Extract text from pages
        for page_num in range(max_pages):
            try:
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            except Exception as e:
                print(f"Warning: Could not extract text from page {page_num + 1}: {str(e)}")
                continue
        
        if not text.strip():
            raise Exception("No text could be extracted from PDF. The file might be scanned or image-based.")
        
        # For very long resumes, truncate to reasonable length (first 50,000 characters)
        if len(text) > 50000:
            print(f"Warning: Resume text is very long ({len(text)} chars). Truncating to 50,000 characters.")
            text = text[:50000]
        
        return text
    except Exception as e:
        raise Exception(f"Error reading PDF: {str(e)}")


def extract_text_from_docx(file):
    """Extract text from DOCX file with better error handling"""
    try:
        doc = docx.Document(file)
        
        # Extract text from paragraphs
        paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
        
        # Extract text from tables
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells if cell.text.strip()]
                if row_text:
                    tables_text.append(' '.join(row_text))
        
        # Combine all text
        text = '\n'.join(paragraphs + tables_text)
        
        if not text.strip():
            raise Exception("No text could be extracted from DOCX file.")
        
        # For very long resumes, truncate to reasonable length (first 50,000 characters)
        if len(text) > 50000:
            print(f"Warning: Resume text is very long ({len(text)} chars). Truncating to 50,000 characters.")
            text = text[:50000]
        
        return text
    except Exception as e:
        raise Exception(f"Error reading DOCX: {str(e)}")


def check_ats_friendliness(text):
    """
    Comprehensive ATS friendliness checker
    Returns: {
        'is_ats_friendly': bool,
        'score': int (0-100),
        'issues': list of issues,
        'suggestions': list of improvements,
        'details': detailed breakdown
    }
    """
    issues = []
    suggestions = []
    score = 100
    details = {}
    
    # Check 1: Minimum length (well-detailed resume)
    if len(text) < 300:
        issues.append("Resume is too short - lacks sufficient detail")
        suggestions.append("Add more details about your experience, skills, and achievements")
        score -= 25
        details['length'] = 'Poor'
    elif len(text) < 800:
        issues.append("Resume could be more detailed")
        suggestions.append("Expand on your key achievements and responsibilities")
        score -= 10
        details['length'] = 'Fair'
    else:
        details['length'] = 'Good'
    
    # Check 2: Contact Information
    text_lower = text.lower()
    
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    phone_pattern = r'(\+\d{1,3}[-.\s]?)?(\(?\d{3}\)?[-.\s]?)?\d{3}[-.\s]?\d{4}'
    
    has_email = bool(re.search(email_pattern, text))
    has_phone = bool(re.search(phone_pattern, text))
    
    if not has_email:
        issues.append("Missing email address")
        suggestions.append("Add a professional email address at the top of your resume")
        score -= 15
    
    if not has_phone:
        issues.append("Missing phone number")
        suggestions.append("Include your contact phone number")
        score -= 10
    
    details['contact_info'] = 'Complete' if (has_email and has_phone) else 'Incomplete'
    
    # Check 3: Essential Sections
    sections_found = []
    sections_missing = []
    
    section_keywords = {
        'Experience': ['experience', 'work history', 'employment', 'professional experience', 'work experience'],
        'Education': ['education', 'qualification', 'degree', 'academic', 'university', 'college'],
        'Skills': ['skills', 'technical skills', 'competencies', 'proficiencies', 'expertise'],
        'Summary': ['summary', 'objective', 'profile', 'about me', 'professional summary']
    }
    
    for section, keywords in section_keywords.items():
        if any(keyword in text_lower for keyword in keywords):
            sections_found.append(section)
        else:
            sections_missing.append(section)
            if section in ['Experience', 'Education', 'Skills']:
                issues.append(f"Missing '{section}' section")
                suggestions.append(f"Add a clear '{section}' section to your resume")
                score -= 15
    
    details['sections'] = f"{len(sections_found)}/4 key sections found"
    
    # Check 4: Action Verbs and Keywords
    action_verbs = ['developed', 'managed', 'led', 'created', 'implemented', 'designed', 
                    'analyzed', 'improved', 'coordinated', 'achieved', 'executed', 
                    'established', 'built', 'optimized', 'delivered', 'increased']
    
    verb_count = sum(1 for verb in action_verbs if verb in text_lower)
    
    if verb_count < 3:
        issues.append("Limited use of strong action verbs")
        suggestions.append("Use more action verbs like: developed, managed, led, implemented, achieved")
        score -= 12
        details['action_verbs'] = 'Poor'
    elif verb_count < 6:
        details['action_verbs'] = 'Fair'
    else:
        details['action_verbs'] = 'Good'
    
    # Check 5: Formatting Issues (special characters from tables/images)
    special_char_ratio = len(re.findall(r'[^\w\s.,;:!?()\-\'/\n]', text)) / max(len(text), 1)
    
    if special_char_ratio > 0.08:
        issues.append("Excessive special characters detected (likely from complex formatting)")
        suggestions.append("Avoid tables, text boxes, and graphics. Use simple bullet points")
        score -= 12
        details['formatting'] = 'Complex (may cause ATS issues)'
    else:
        details['formatting'] = 'Simple (ATS-friendly)'
    
    # Check 6: Bullet Points
    bullet_indicators = ['•', '◦', '○', '■', '▪', '-', '*', '►', '→']
    has_bullets = any(char in text for char in bullet_indicators)
    
    if not has_bullets and len(text) > 500:
        issues.append("No bullet points found - content may be hard to parse")
        suggestions.append("Use bullet points to list your responsibilities and achievements")
        score -= 8
    
    # Check 7: Dates (for experience/education)
    date_patterns = [
        r'\b(19|20)\d{2}\b',  # Years
        r'\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]* \d{4}\b',  # Month Year
        r'\d{1,2}/\d{4}',  # MM/YYYY
    ]
    
    date_count = sum(len(re.findall(pattern, text)) for pattern in date_patterns)
    
    if date_count < 2:
        issues.append("Missing dates for experience or education")
        suggestions.append("Include dates (MM/YYYY format) for your work experience and education")
        score -= 8
    
    # Check 8: Length of experience descriptions
    if 'experience' in text_lower:
        exp_index = text_lower.find('experience')
        content_after_exp = text[exp_index:exp_index+500] if exp_index != -1 else ""
        
        if len(content_after_exp) < 200:
            issues.append("Work experience section seems too brief")
            suggestions.append("Provide more details about your roles, responsibilities, and achievements")
            score -= 10
    
    # Check 9: Headers/Titles
    caps_words = re.findall(r'\b[A-Z]{2,}\b', text)
    if len(caps_words) < 3 and len(text) > 500:
        suggestions.append("Consider using clear section headers (e.g., EXPERIENCE, EDUCATION, SKILLS)")
        score -= 5
    
    # Final score adjustment
    score = max(0, min(100, score))
    
    # Determine if ATS-friendly (threshold: 70)
    is_ats_friendly = score >= 70
    
    # Overall assessment
    if score >= 85:
        overall = "Excellent - Highly ATS-friendly"
    elif score >= 70:
        overall = "Good - ATS-friendly with minor improvements possible"
    elif score >= 50:
        overall = "Fair - Needs improvement for better ATS compatibility"
    else:
        overall = "Poor - Major improvements needed"
    
    return {
        'is_ats_friendly': is_ats_friendly,
        'score': score,
        'overall': overall,
        'issues': issues,
        'suggestions': suggestions,
        'details': details
    }


@app.route('/api/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({'status': 'ok', 'message': 'API is running'})


@app.route('/api/upload-resume', methods=['POST'])
def upload_resume():
    """
    Upload resume file, check ATS friendliness, and analyze if friendly
    """
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400
        
        file = request.files['file']
        
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        filename = file.filename.lower()
        
        # Check file extension
        if not (filename.endswith('.pdf') or filename.endswith('.docx')):
            return jsonify({'error': 'Invalid file format. Please upload PDF or DOCX'}), 400
        
        # Extract text based on file type
        try:
            if filename.endswith('.pdf'):
                resume_text = extract_text_from_pdf(file)
            else:  # .docx
                resume_text = extract_text_from_docx(file)
        except Exception as e:
            return jsonify({'error': str(e)}), 400
        
        if not resume_text or len(resume_text.strip()) < 50:
            return jsonify({'error': 'Could not extract sufficient text from file. Please ensure the file contains readable text.'}), 400
        
        # Check ATS friendliness
        ats_result = check_ats_friendliness(resume_text)
        
        response = {
            'ats_check': ats_result,
            'resume_text_length': len(resume_text)
        }
        
        # If ATS-friendly, analyze with ML model
        if ats_result['is_ats_friendly']:
            try:
                prediction = predictor.predict(resume_text)
                predicted_role = prediction['predicted_role']
                questions = interview_questions.get(predicted_role, interview_questions['DEFAULT'])
                
                response['analysis'] = {
                    'predicted_role': predicted_role,
                    'confidence': prediction['confidence'],
                    'top_3_roles': prediction['top_3_roles'],
                    'interview_questions': questions
                }
            except Exception as e:
                print(f"Analysis error: {str(e)}")
                response['analysis_error'] = f"Could not analyze resume: {str(e)}"
        
        return jsonify(response), 200
        
    except Exception as e:
        print(f"Server error: {str(e)}")
        return jsonify({'error': f"Server error: {str(e)}"}), 500


@app.route('/api/analyze-resume', methods=['POST'])
def analyze_resume():
    """Analyze resume text directly (for resume builder integration)"""
    try:
        data = request.get_json()
        resume_text = data.get('resume_text', '')
        
        if not resume_text:
            return jsonify({'error': 'No resume text provided'}), 400
        
        # Check ATS first
        ats_result = check_ats_friendliness(resume_text)
        
        response = {
            'ats_check': ats_result
        }
        
        # Analyze if ATS-friendly
        if ats_result['is_ats_friendly']:
            try:
                prediction = predictor.predict(resume_text)
                predicted_role = prediction['predicted_role']
                questions = interview_questions.get(predicted_role, interview_questions['DEFAULT'])
                
                response['analysis'] = {
                    'predicted_role': predicted_role,
                    'confidence': prediction['confidence'],
                    'top_3_roles': prediction['top_3_roles'],
                    'interview_questions': questions
                }
            except Exception as e:
                response['analysis_error'] = f"Could not analyze resume: {str(e)}"
        
        return jsonify(response), 200
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.errorhandler(413)
def request_entity_too_large(error):
    """Handle file too large error"""
    return jsonify({'error': 'File too large. Maximum size is 50MB'}), 413


if __name__ == '__main__':
    print("="*60)
    print("🚀 Resume Analysis API with ATS Checker")
    print("="*60)
    print("📍 API running at: http://localhost:5000")
    print("\n🔗 Available endpoints:")
    print("  GET  /api/health")
    print("  POST /api/upload-resume")
    print("  POST /api/analyze-resume")
    print("\n" + "="*60)
    app.run(debug=True, port=5000, host='0.0.0.0')